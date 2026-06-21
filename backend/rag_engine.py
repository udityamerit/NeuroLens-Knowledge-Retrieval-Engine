import os
import pypdf
import docx
from typing import List, Dict, Any, Optional
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from dotenv import load_dotenv

# Suppress some HF warnings
os.environ["TOKENIZERS_PARALLELISM"] = "false"

# Load environment variables from projects root .env
dotenv_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), ".env")
load_dotenv(dotenv_path=dotenv_path)

class RAGEngine:
    def __init__(self, index_dir: str = "faiss_index"):
        self.index_dir = index_dir
        self.embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
        self.vector_store = None
        self.uploaded_files = [] # Stores list of file names that have been indexed
        
        # Load index if it exists
        if os.path.exists(self.index_dir) and os.path.isdir(self.index_dir):
            try:
                self.vector_store = FAISS.load_local(
                    self.index_dir, 
                    self.embeddings, 
                    allow_dangerous_deserialization=True
                )
                # Read list of files from vector store if possible
                if self.vector_store:
                    # Collect all unique sources from docstore
                    sources = set()
                    for doc_id, doc in self.vector_store.docstore._dict.items():
                        source = doc.metadata.get("source", "Unknown")
                        sources.add(source)
                    self.uploaded_files = list(sources)
            except Exception as e:
                print(f"Error loading index: {e}. Clearing corrupted index directory.")
                self.vector_store = None
                try:
                    import shutil
                    shutil.rmtree(self.index_dir)
                except Exception:
                    pass

    def load_document(self, file_path: str, filename: str) -> List[Document]:
        """Loads a document and returns a list of Document objects with metadata."""
        documents = []
        ext = os.path.splitext(file_path)[1].lower()
        print(f"📖 RAGEngine: Loading file '{filename}' from temporary path '{file_path}' (extension: {ext})")
        
        if not os.path.exists(file_path):
            print(f"  ❌ File does not exist at path: {file_path}")
            return []
            
        file_size = os.path.getsize(file_path)
        print(f"  File size: {file_size} bytes")
        
        if file_size == 0:
            print("  ⚠️ File is empty (0 bytes). Skipping.")
            return []

        if ext == ".pdf":
            try:
                with open(file_path, "rb") as f:
                    reader = pypdf.PdfReader(f)
                    for page_idx, page in enumerate(reader.pages):
                        text = page.extract_text()
                        if text and text.strip():
                            documents.append(Document(
                                page_content=text,
                                metadata={
                                    "source": filename,
                                    "page": page_idx + 1,
                                    "type": "pdf"
                                }
                            ))
                print(f"  Processed PDF: {len(reader.pages)} pages, extracted {len(documents)} document pages.")
            except Exception as e:
                import traceback
                print(f"  ❌ Error reading PDF {filename}:")
                traceback.print_exc()
                
        elif ext in [".docx", ".doc"]:
            try:
                with open(file_path, "rb") as f:
                    doc = docx.Document(f)
                    # Read paragraph by paragraph
                    paragraphs_text = []
                    for p in doc.paragraphs:
                        if p.text and p.text.strip():
                            paragraphs_text.append(p.text)
                    
                    # Combine paragraphs but keep them structured
                    full_text = "\n\n".join(paragraphs_text)
                    if full_text.strip():
                        documents.append(Document(
                            page_content=full_text,
                            metadata={
                                "source": filename,
                                "type": "docx"
                            }
                        ))
                print(f"  Processed DOCX: extracted {len(documents)} document blocks.")
            except Exception as e:
                import traceback
                print(f"  ❌ Error reading DOCX {filename}:")
                traceback.print_exc()
                
        elif ext in [".png", ".jpg", ".jpeg", ".webp", ".gif"]:
            try:
                import base64
                import json
                import urllib.request
                
                print(f"  🖼️ Image file detected: {filename}. Processing with Vision LLM...")
                
                # Check for API keys in environment
                groq_key = os.getenv("GROQ_API_KEY")
                openai_key = os.getenv("OPENAI_API_KEY")
                
                with open(file_path, "rb") as img_file:
                    img_data = img_file.read()
                    base64_str = base64.b64encode(img_data).decode("utf-8")
                
                # Determine mime type
                mime_type = "image/png"
                if ext == ".jpg" or ext == ".jpeg":
                    mime_type = "image/jpeg"
                elif ext == ".webp":
                    mime_type = "image/webp"
                elif ext == ".gif":
                    mime_type = "image/gif"
                
                extracted_text = None
                
                # Try Groq first if key exists
                if groq_key:
                    url = "https://api.groq.com/openai/v1/chat/completions"
                    headers = {
                        "Authorization": f"Bearer {groq_key}",
                        "Content-Type": "application/json",
                        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
                    }
                    groq_models = [
                        "meta-llama/llama-4-scout-17b-16e-instruct",
                        "qwen/qwen3.6-27b"
                    ]
                    
                    last_err = None
                    for model in groq_models:
                        try:
                            print(f"    Trying Groq model: {model}...")
                            payload = {
                                "model": model,
                                "messages": [
                                    {
                                        "role": "user",
                                        "content": [
                                            {
                                                "type": "text",
                                                "text": "You are an advanced document analyst and OCR engine. Describe this image in detail, transcribing all text, labels, structures, charts, graphs, or tables word-for-word. Provide a clear, structured textual description without repeating sentences or phrases. Avoid looping or duplicating descriptive statements."
                                            },
                                            {
                                                "type": "image_url",
                                                "image_url": {
                                                    "url": f"data:{mime_type};base64,{base64_str}"
                                                }
                                            }
                                        ]
                                    }
                                ],
                                "temperature": 0.5
                            }
                            req = urllib.request.Request(
                                url,
                                data=json.dumps(payload).encode("utf-8"),
                                headers=headers,
                                method="POST"
                            )
                            with urllib.request.urlopen(req) as res:
                                res_data = json.loads(res.read().decode("utf-8"))
                                extracted_text = res_data["choices"][0]["message"]["content"]
                            break # Success! Break out of model loop
                        except Exception as e:
                            print(f"      Model {model} failed: {e}")
                            last_err = e
                    
                    if not extracted_text and not openai_key:
                        raise last_err or Exception("All Groq vision models failed.")
                
                # Try OpenAI if Groq key doesn't exist but OpenAI key does
                elif openai_key:
                    print("    Using OpenAI GPT-4o-mini Vision Model...")
                    url = "https://api.openai.com/v1/chat/completions"
                    headers = {
                        "Authorization": f"Bearer {openai_key}",
                        "Content-Type": "application/json"
                    }
                    payload = {
                        "model": "gpt-4o-mini",
                        "messages": [
                            {
                                "role": "user",
                                "content": [
                                    {
                                        "type": "text",
                                        "text": "You are an advanced document analyst and OCR engine. Describe this image in extreme detail, transcribing all text, labels, structures, charts, graphs, or tables word-for-word. Provide a complete textual description so it can be fully indexable by a search engine."
                                    },
                                    {
                                        "type": "image_url",
                                        "image_url": {
                                            "url": f"data:{mime_type};base64,{base64_str}"
                                        }
                                    }
                                ]
                            }
                        ],
                        "temperature": 0.1
                    }
                    req = urllib.request.Request(
                        url,
                        data=json.dumps(payload).encode("utf-8"),
                        headers=headers,
                        method="POST"
                    )
                    with urllib.request.urlopen(req) as res:
                        res_data = json.loads(res.read().decode("utf-8"))
                        extracted_text = res_data["choices"][0]["message"]["content"]
                
                else:
                    raise Exception("No vision-capable API keys.")
                
                if extracted_text:
                    documents.append(Document(
                        page_content=extracted_text,
                        metadata={
                            "source": filename,
                            "type": ext[1:]
                        }
                    ))
                    print(f"  Processed Image: extracted {len(documents)} document blocks.")
                    
            except Exception as e:
                import traceback
                print(f"  ❌ Error processing image {filename}:")
                traceback.print_exc()

        else: # Default to plain text
            try:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
                    if text.strip():
                        documents.append(Document(
                            page_content=text,
                            metadata={
                                "source": filename,
                                "type": "txt"
                            }
                        ))
                print(f"  Processed Text File: extracted {len(documents)} document blocks.")
            except Exception as e:
                import traceback
                print(f"  ❌ Error reading text file {filename}:")
                traceback.print_exc()

        return documents

    def index_files(self, file_paths: List[str], filenames: List[str]) -> bool:
        """Processes and indexes a list of files into the local FAISS database."""
        all_documents = []
        for path, name in zip(file_paths, filenames):
            docs = self.load_document(path, name)
            all_documents.extend(docs)

        if not all_documents:
            return False

        # Split documents into chunks (tuned parameters for higher accuracy RAG retrieval)
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=150)
        chunks = text_splitter.split_documents(all_documents)

        # Add chunks to vector store
        if self.vector_store is None:
            self.vector_store = FAISS.from_documents(chunks, self.embeddings)
        else:
            self.vector_store.add_documents(chunks)

        # Save index locally
        os.makedirs(self.index_dir, exist_ok=True)
        self.vector_store.save_local(self.index_dir)
        
        # Update file list
        for name in filenames:
            if name not in self.uploaded_files:
                self.uploaded_files.append(name)
        
        return True

    def get_relevant_chunks(self, query: str, k: int = 5) -> List[Dict[str, Any]]:
        """Retrieves top k matching text chunks from the vector store."""
        if self.vector_store is None:
            return []

        results = self.vector_store.similarity_search_with_score(query, k=k)
        chunks = []
        for doc, score in results:
            chunks.append({
                "content": doc.page_content,
                "metadata": doc.metadata,
                "relevance_score": float(score)
            })
        return chunks

    def clear(self):
        """Clears the current vector store index and document registry."""
        self.vector_store = None
        self.uploaded_files = []
        if os.path.exists(self.index_dir):
            for file in os.listdir(self.index_dir):
                os.remove(os.path.join(self.index_dir, file))
            try:
                os.rmdir(self.index_dir)
            except Exception:
                pass

    def generate_standalone_query(
        self,
        query: str,
        history: List[Dict[str, Any]],
        provider: str,
        api_key: str,
        model_name: str
    ) -> str:
        """Uses the LLM to rewrite the user query as a standalone search query incorporating context from history."""
        if not history:
            return query
            
        history_msgs = [m for m in history if m.get("role") in ["user", "assistant"] and m.get("content")]
        if not history_msgs:
            return query

        # Limit history to the last 4 messages to prevent excessive token count or latency
        recent_history = history_msgs[-4:]
        
        # Build rephrasing prompt
        history_str = ""
        for msg in recent_history:
            role = "User" if msg["role"] == "user" else "Assistant"
            content = msg["content"]
            # Strip very long outputs
            if len(content) > 300:
                content = content[:300] + "..."
            history_str += f"{role}: {content}\n"

        rephrase_prompt = (
            "Given the following chat history and a follow-up question, "
            "rephrase the follow-up question to be a self-contained standalone search query. "
            "Do NOT answer the question. Just rephrase it to include necessary details "
            "from the history so that it can be searched in a vector database.\n"
            "Return ONLY the raw standalone question string and absolutely nothing else.\n\n"
            f"Chat History:\n{history_str}\n"
            f"Follow-up Question: {query}\n"
            "Standalone Question:"
        )

        try:
            if provider.lower() == "groq":
                from langchain_groq import ChatGroq
                llm = ChatGroq(
                    api_key=api_key,
                    model_name=model_name,
                    temperature=0.0
                )
                response = llm.invoke([
                    {"role": "user", "content": rephrase_prompt}
                ])
                rephrased = response.content.strip().strip('"').strip("'")
                print(f"🔄 Standalone Query generated: '{rephrased}' (original: '{query}')")
                return rephrased if rephrased else query

            elif provider.lower() == "openai":
                from langchain_community.chat_models import ChatOpenAI
                llm = ChatOpenAI(
                    api_key=api_key,
                    model_name=model_name,
                    temperature=0.0
                )
                response = llm.invoke([
                    {"role": "user", "content": rephrase_prompt}
                ])
                rephrased = response.content.strip().strip('"').strip("'")
                print(f"🔄 Standalone Query generated: '{rephrased}' (original: '{query}')")
                return rephrased if rephrased else query

        except Exception as e:
            print(f"⚠️ Failed to generate standalone query: {e}. Using raw query.")
            
        return query

    def query_llm(
        self,
        query: str,
        history: List[Dict[str, Any]],
        provider: str,
        api_key: str,
        model_name: str,
        temperature: float = 0.3,
        k: int = 5
    ) -> Dict[str, Any]:
        """Runs the complete RAG pipeline: retrieves chunks and queries the LLM."""
        # 1. Resolve API key: use incoming key first, otherwise fallback to system/dotenv variables
        resolved_key = api_key.strip() if api_key else ""
        if not resolved_key:
            if provider.lower() == "groq":
                resolved_key = os.environ.get("GROQ_API_KEY", "")
            elif provider.lower() == "openai":
                resolved_key = os.environ.get("OPENAI_API_KEY", "")
            elif provider.lower() == "huggingface":
                resolved_key = os.environ.get("HF_TOKEN", "") or os.environ.get("HUGGINGFACEHUB_API_TOKEN", "")

        # 2. Generate standalone search query using chat history
        search_query = query
        if history and resolved_key:
            search_query = self.generate_standalone_query(query, history, provider, resolved_key, model_name)

        # 3. Retrieve documents
        relevant_chunks = self.get_relevant_chunks(search_query, k=k)
        
        # Check if the user is asking about the creator/builder/developer of the project or studying the author
        creator_keywords = ["creator", "create", "build", "developer", "who made", "who built", "who programmed", "author", "owner", "tiwari", "uditya", "programmer", "education", "experience", "resume", "cv", "portfolio", "study"]
        author_study_keywords = ["author study", "study the author", "study the creator", "author's projects", "creator's projects", "github projects"]
        
        is_asking_about_creator = any(k in query.lower() for k in creator_keywords)
        is_asking_about_author_study = any(k in query.lower() for k in author_study_keywords)

        # Developer Bio based on his actual resume
        developer_bio = (
            "You were created and developed by Uditya Narayan Tiwari, a highly skilled Machine Learning Engineer and Generative AI Developer.\n"
            "Key credentials & background of Uditya:\n"
            "- **Education**: Currently pursuing B.Tech in Computer Science Engineering with specialization in AI & ML at VIT Bhopal University (2023 - Present) with an impressive CGPA of 9.00.\n"
            "- **Experience**: Technical Member & Community Admin of the Microsoft Learn Student Chapter (2025 - 2026), and Core Tech Team Member of the Matrix Media Club at VIT Bhopal (Jan 2025 - Jul 2025).\n"
            "- **Key Projects**:\n"
            "  1. **NeuroLens**: This exact AI-powered Knowledge Retrieval Engine (Python, LangChain, FAISS, RAG, LLMs).\n"
            "  2. **Smart Medical Care For Rural Areas**: An AI symptom-to-medicine recommendation model using NLP, Sentence Transformers (cosine similarity + MMR), and Flask.\n"
            "  3. **Breast Cancer Classification**: Machine learning model comparing SVM, Logistic Regression, Random Forest, achieving ~97% accuracy.\n"
            "- **Certifications & Awards**: Geodata Processing & AI/ML Geodata Analysis certifications from the Indian Institute of Remote Sensing (IIRS), ISRO (2024); First Prize Winner in the VIT Bhopal Hackathon (Feb 2025); Final Round Qualifier in the JHU Hackathon (Jan 2025).\n\n"
            "You MUST present the links in standard Markdown format so they render as clean clickable links:\n"
            "- [Personal Portfolio](https://udityanarayantiwari.netlify.app/)\n"
            "- [GitHub Profile](https://github.com/udityamerit)\n"
            "- [LinkedIn Profile](https://www.linkedin.com/in/uditya-narayan-tiwari-562332289/)\n"
            "- [Knowledge Base](https://udityaknowledgebase.netlify.app/)\n\n"
        )

        if not relevant_chunks:
            if is_asking_about_author_study or is_asking_about_creator:
                system_prompt = (
                    "You are NeuroLens, an advanced AI document intelligence engine. "
                    "When asked about your creator, developer, programmer, or builder, or asked to study your author/projects, you must answer with his real resume profile:\n\n"
                    f"{developer_bio}"
                    "Respond with extreme professionalism and pride in Uditya's engineering work."
                )
            else:
                system_prompt = (
                    "You are NeuroLens, an advanced AI assistant. "
                    "Respond to the user's question helpfully and clearly. "
                    "Respond in the same language as the user's question."
                )
        else:
            # 4. Build prompt context
            context_str = ""
            for i, chunk in enumerate(relevant_chunks):
                source = chunk["metadata"].get("source", "Unknown")
                page_info = f" (Page {chunk['metadata'].get('page')})" if "page" in chunk["metadata"] else ""
                context_str += f"--- Source {i+1}: {source}{page_info} ---\n{chunk['content']}\n\n"

            if is_asking_about_author_study or is_asking_about_creator:
                system_prompt = (
                    "You are NeuroLens, an advanced AI document intelligence engine. "
                    "In addition to answering from the documents, when asked about your creator, developer, programmer, builder, or asked to study your author/projects, you must respond with his real resume profile:\n\n"
                    f"{developer_bio}"
                    "Explain that you are analyzing the documents loaded into your library, but first proudly introduce Uditya Narayan Tiwari as your creator.\n\n"
                    f"Here is the context from the documents:\n\n{context_str}"
                )
            else:
                system_prompt = (
                    "You are NeuroLens, an advanced AI document analyst. "
                    "Your task is to answer the user's question based strictly on the provided context source blocks. "
                    "Respond in the same language as the user's question (e.g., if the user asks in Hindi, translate the relevant context facts and answer in Hindi). "
                    "For each statement you make, try to cite which Source (e.g., [Source 1], [Source 2]) you retrieved the information from. "
                    "If the context does not contain the information needed to answer the question, state that you cannot find the answer in the provided documents.\n\n"
                    f"Here is the context retrieved from the documents:\n\n{context_str}"
                )

        # 5. Call LLM
        answer = ""
        try:
            if provider.lower() == "groq":
                if not resolved_key:
                    raise ValueError("Groq API key not found. Please provide it in Settings.")
                from langchain_groq import ChatGroq
                llm = ChatGroq(
                    api_key=resolved_key,
                    model_name=model_name,
                    temperature=temperature
                )
                
                # Build messages list incorporating chat history
                messages = [{"role": "system", "content": system_prompt}]
                if history:
                    for msg in history:
                        role = msg.get("role")
                        content = msg.get("content")
                        if role in ["user", "assistant"] and content and not content.startswith("📥 **System:**") and not content.startswith("❌ **Error"):
                            messages.append({"role": role, "content": content})
                messages.append({"role": "user", "content": query})
                
                response = llm.invoke(messages)
                answer = response.content

            elif provider.lower() == "openai":
                if not resolved_key:
                    raise ValueError("OpenAI API key not found. Please provide it in Settings.")
                from langchain_community.chat_models import ChatOpenAI
                llm = ChatOpenAI(
                    api_key=resolved_key,
                    model_name=model_name,
                    temperature=temperature
                )
                
                # Build messages list incorporating chat history
                messages = [{"role": "system", "content": system_prompt}]
                if history:
                    for msg in history:
                        role = msg.get("role")
                        content = msg.get("content")
                        if role in ["user", "assistant"] and content and not content.startswith("📥 **System:**") and not content.startswith("❌ **Error"):
                            messages.append({"role": role, "content": content})
                messages.append({"role": "user", "content": query})
                
                response = llm.invoke(messages)
                answer = response.content

            elif provider.lower() == "huggingface":
                if not resolved_key:
                    raise ValueError("Hugging Face API token not found. Please provide it in Settings.")
                from langchain_huggingface import HuggingFaceEndpoint
                llm = HuggingFaceEndpoint(
                    repo_id=model_name,
                    huggingfacehub_api_token=resolved_key,
                    temperature=temperature,
                    timeout=30
                )
                
                # Construct formatted text containing prompt history
                full_prompt = f"System: {system_prompt}\n"
                if history:
                    for msg in history:
                        role = msg.get("role")
                        content = msg.get("content")
                        if role in ["user", "assistant"] and content and not content.startswith("📥 **System:**") and not content.startswith("❌ **Error"):
                            speaker = "User" if role == "user" else "Assistant"
                            full_prompt += f"{speaker}: {content}\n"
                full_prompt += f"User: {query}\nAssistant:"
                
                answer = llm.invoke(full_prompt)

            else:
                # Fallback mock/rule-based answer if no valid provider
                answer = (
                    "⚠️ NeuroLens: LLM Provider not recognized or API Key not configured. "
                    "However, here is the text retrieved from your documents that matches your query:\n\n"
                )
                for chunk in relevant_chunks[:2]:
                    source = chunk["metadata"].get("source", "Unknown")
                    page_info = f" (Page {chunk['metadata'].get('page')})" if "page" in chunk["metadata"] else ""
                    answer += f"📄 **From {source}{page_info}:**\n> {chunk['content']}\n\n"
                answer += "\nConfigure a valid LLM API Key (Groq, OpenAI, or Hugging Face) in the settings to generate synthesis answers."

        except Exception as e:
            answer = f"Error calling {provider} LLM: {str(e)}\n\n"
            answer += "⚠️ Please verify your API Key and Model configuration in the Settings panel."

        return {
            "answer": answer,
            "sources": [
                {
                    "content": chunk["content"],
                    "source": chunk["metadata"].get("source", "Unknown"),
                    "page": chunk["metadata"].get("page"),
                    "type": chunk["metadata"].get("type", "txt")
                }
                for chunk in relevant_chunks
            ]
        }
